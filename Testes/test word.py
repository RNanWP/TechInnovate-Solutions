from docx import Document
from docx.shared import Pt
import os

# Defina o caminho para salvar o documento na área de trabalho
desktop = os.path.join(os.path.join(os.environ['USERPROFILE']), 'Desktop')
file_path = os.path.join(desktop, 'TechInnovate_Project_Documentation.docx')

# Criação do documento
doc = Document()

# Capa
doc.add_heading('TechInnovate Solutions', level=1)
doc.add_paragraph('Nome da Empresa: TechInnovate Solutions')
doc.add_paragraph('Ramo da Empresa: Tecnologia')
doc.add_paragraph('Link Github: [link]')
doc.add_paragraph('Equipe Responsável: [Nome dos Membros da Equipe e RA]')
doc.add_paragraph('Curso: [Curso]')
doc.add_paragraph('Turma: 40')
doc.add_paragraph('Semestre: 1')
doc.add_paragraph('Ano: 2024')

doc.add_page_break()

# Índice
doc.add_heading('Índice', level=1)

doc.add_paragraph('1. Escopo …………………………………………………………… 2')
doc.add_paragraph('2. Serviços ………………………………………………………….. 3')
doc.add_paragraph('3. Estruturação Interna da Empresa …………………………… 4')
doc.add_paragraph('  3.1 Aprendizado de Máquina …………………………………… 4')
doc.add_paragraph('    3.1.1 Entrega 1 …………………………………………………… 4')
doc.add_paragraph('    3.1.2 Entrega 2 …………………………………………………… 5')
doc.add_paragraph('  3.2 Ciência de Dados ……………………………………………… 5')
doc.add_paragraph('    3.2.1 Entrega 1 …………………………………………………… 5')
doc.add_paragraph('    3.2.2 Entrega 2 …………………………………………………… 6')
doc.add_paragraph('  3.3 Modelagem de Dados ………………………………………… 6')
doc.add_paragraph('  3.4 Redes de Computadores …………………………………… 6')
doc.add_paragraph('  3.5 Segurança da Informação ………………………………… 6')

doc.add_page_break()

# Escopo do Projeto
doc.add_heading('1. Escopo do Projeto', level=1)
doc.add_paragraph('O presente documento detalha o escopo do projeto proposto pela empresa TechInnovate Solutions. O projeto tem como objetivo [descrição sucinta do objetivo do projeto], visando [benefícios esperados ou problemas a serem resolvidos].')

doc.add_page_break()

# Serviços Oferecidos
doc.add_heading('2. Serviços Oferecidos', level=1)
doc.add_paragraph('A empresa TechInnovate Solutions oferecerá os seguintes serviços:')
doc.add_paragraph('- Desenvolvimento de Software')
doc.add_paragraph('- Consultoria em Tecnologia')
doc.add_paragraph('- Suporte Técnico')

doc.add_page_break()

# Estruturação Interna da Empresa
doc.add_heading('3. Estruturação Interna da Empresa', level=1)

# Aprendizado de Máquina
doc.add_heading('3.1 Aprendizado de Máquina', level=2)

# Entrega 1
doc.add_heading('3.1.1 Entrega 1', level=3)
doc.add_paragraph('Implementação de um modelo de previsão de demanda utilizando Random Forest e otimização de hiperparâmetros com GridSearchCV.')

# Detalhamento da Entrega 1
doc.add_heading('Documentação do Processo de Construção e Treinamento do Modelo', level=4)
doc.add_paragraph('''Introdução:
Este documento fornece uma visão detalhada do processo de construção e treinamento do modelo de aprendizado de máquina para a tarefa específica. Descreve as etapas, parâmetros selecionados e os resultados obtidos durante o desenvolvimento do modelo.

Objetivo:
O objetivo principal deste modelo é prever a demanda futura de produtos para otimizar o estoque da empresa.

Etapas do Processo:
1. Exploração de Dados e Pré-processamento:
   Coleta de Dados:
       Descreve as fontes de dados utilizadas.
       Lista de variáveis/features incluídas.
   Limpeza e Pré-processamento:
       Identificação e tratamento de valores ausentes, outliers, etc.
       Transformações aplicadas aos dados.
2. Implementação de Modelos de Aprendizado de Máquina:
   Escolha de Algoritmos:
       Justificativa para a escolha dos algoritmos utilizados.
   Implementação:
       Detalhes sobre como os modelos foram implementados.
       Utilização de bibliotecas (por exemplo, Scikit-learn, TensorFlow).
3. Otimização e Validação do Modelo:
   Otimização de Hiperparâmetros:
       Descrição do processo de otimização.
       Lista dos hiperparâmetros ajustados.
   Validação Cruzada:
       Detalhes sobre como a validação cruzada foi realizada.
       Resultados obtidos.
   Parâmetros do Modelo:
       Lista completa de hiperparâmetros e seus valores finais após a otimização.
       Outros parâmetros relevantes para o modelo.
   Métricas de Avaliação:
       Descrição das métricas utilizadas para avaliar o desempenho do modelo.
       Resultados específicos obtidos para cada métrica.
''')

# Entrega 2
doc.add_heading('3.1.2 Entrega 2', level=3)
doc.add_paragraph('Implementação de dados aleatórios e visualização de tabelas utilizando Faker e Tkinter.')

# Ciência de Dados
doc.add_heading('3.2 Ciência de Dados', level=2)

# Entrega 1
doc.add_heading('3.2.1 Entrega 1', level=3)
doc.add_paragraph('Análise exploratória de dados e visualização de dados utilizando pandas e matplotlib.')

# Entrega 2
doc.add_heading('3.2.2 Entrega 2', level=3)
doc.add_paragraph('Modelagem preditiva e avaliação de desempenho do modelo utilizando Scikit-learn.')

# Modelagem de Dados
doc.add_heading('3.3 Modelagem de Dados', level=2)
doc.add_paragraph('Criação de um diagrama de entidade-relacionamento e normalização do banco de dados para a empresa TechInnovate Solutions.')

# Redes de Computadores
doc.add_heading('3.4 Redes de Computadores', level=2)
doc.add_paragraph('Configuração de rede local e política de segurança para garantir a integridade e confidencialidade dos dados.')

# Segurança da Informação
doc.add_heading('3.5 Segurança da Informação', level=2)
doc.add_paragraph('Implementação de medidas de segurança para proteger os dados sensíveis da empresa contra acessos não autorizados e ataques cibernéticos.')

# Salvando o documento
doc.save(file_path)
